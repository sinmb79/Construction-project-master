from __future__ import annotations

from pathlib import Path

from docx import Document

from civilplan_mcp.config import get_settings
from civilplan_mcp.models import ProjectDomain
from civilplan_mcp.tools._base import wrap_response


def generate_budget_report(
    *,
    report_type: str,
    project_data: dict,
    boq_summary: dict,
    department: str,
    requester: str,
    output_filename: str | None = None,
) -> dict:
    output_dir = Path(project_data.get("output_dir", get_settings().output_dir))
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / (output_filename or f"{report_type}.docx")

    document = Document()
    document.add_heading(report_type, level=0)
    document.add_paragraph(f"Department: {department}")
    document.add_paragraph(f"Requester: {requester}")
    document.add_paragraph(f"Project domain: {project_data.get('domain')}")
    document.add_paragraph(f"Total cost: {boq_summary['total_cost']:,} KRW")
    document.add_paragraph("Conceptual planning document only. Not for official submission.")
    document.save(path)

    return wrap_response({"status": "success", "file_path": str(path)}, ProjectDomain.복합)
