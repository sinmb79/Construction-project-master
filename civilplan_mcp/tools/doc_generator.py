from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from civilplan_mcp.config import get_settings
from civilplan_mcp.models import ProjectDomain
from civilplan_mcp.tools._base import wrap_response


def generate_investment_doc(
    *,
    project_name: str,
    project_spec: dict[str, Any],
    quantities: dict[str, Any],
    legal_procedures: dict[str, Any],
    boq_summary: dict[str, Any],
    requester: str = "22B Labs",
    output_filename: str | None = None,
) -> dict[str, Any]:
    output_dir = Path(project_spec.get("output_dir", get_settings().output_dir))
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / (output_filename or f"{project_name}_investment.docx")

    document = Document()
    document.add_heading(project_name, level=0)
    document.add_paragraph(f"Requester: {requester}")
    document.add_paragraph("This document is for conceptual planning only and is not valid for official submission.")

    document.add_heading("1. Project Overview", level=1)
    document.add_paragraph(f"Domain: {project_spec.get('domain')}")
    document.add_paragraph(f"Region: {project_spec.get('region')}")
    document.add_paragraph(f"Period: {project_spec.get('year_start')} ~ {project_spec.get('year_end')}")

    document.add_heading("2. Quantity Summary", level=1)
    for category, items in quantities["quantities"].items():
        document.add_paragraph(f"{category}: {', '.join(f'{k}={v}' for k, v in items.items())}")

    document.add_heading("3. Cost Summary", level=1)
    document.add_paragraph(f"Direct cost: {boq_summary['direct_cost']:,} KRW")
    document.add_paragraph(f"Indirect cost: {boq_summary['indirect_cost']:,} KRW")
    document.add_paragraph(f"Total cost: {boq_summary['total_cost']:,} KRW")

    document.add_heading("4. Procedures", level=1)
    document.add_paragraph(f"Procedure count: {legal_procedures['summary']['total_procedures']}")
    for phase, procedures in legal_procedures.get("phases", {}).items():
        document.add_paragraph(f"{phase}: {len(procedures)} item(s)")

    document.save(path)

    return wrap_response({"status": "success", "file_path": str(path)}, ProjectDomain.복합)
